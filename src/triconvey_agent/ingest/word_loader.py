"""Word document loader (.docx / .doc).

Extracts text from Microsoft Word files and returns a ``Document`` object
using the same schema as the PDF loader.  Text is extracted paragraph-by-
paragraph; each paragraph is treated as plain text and the full document is
joined into ``raw_text``.  Document type classification re-uses the same
signature-scoring logic as the PDF loader.

Requires: ``python-docx``  (pip install python-docx)
"""
from __future__ import annotations

from pathlib import Path

from triconvey_agent.ingest.normalizer import (
    collapse_snippet,
    normalize_text,
)
from triconvey_agent.ingest.pdf_loader import classify_pdf_document
from triconvey_agent.schemas.documents import (
    Document,
    DocumentPage,
    InputFileType,
)

# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_docx_text(path: Path) -> tuple[str, list[str]]:
    """Return (full_raw_text, list_of_section_texts) for a .docx file.

    Sections are split on Heading-style paragraphs so we can expose
    DocumentPage objects that correspond to logical sections rather than
    physical pages (Word has no reliable page-break API without Win32).
    """
    try:
        import docx  # type: ignore[import]
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "python-docx is not installed. "
            "Run:  pip install python-docx"
        ) from exc

    doc = docx.Document(str(path))

    sections: list[list[str]] = [[]]   # at least one section
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Split into a new section on any Heading style
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading") and sections[-1]:
            sections.append([])
        sections[-1].append(text)

    # Also pull text from tables
    table_lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_lines.append("  ".join(row_cells))

    if table_lines:
        sections.append(table_lines)

    section_texts = ["\n".join(lines) for lines in sections if lines]
    full_text = "\n\n".join(section_texts)
    return full_text, section_texts


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def load_word_document(path: str | Path) -> Document:
    """Load a .docx (or .doc) file and return a ``Document``.

    .doc (legacy binary) requires ``python-docx`` ≥ 1.x which reads .docx
    natively; true binary .doc is not supported — the file must be re-saved
    as .docx first.  A clear error is raised if the format is unreadable.
    """
    file_path = Path(path).expanduser().resolve()

    raw_text, section_texts = _extract_docx_text(file_path)
    normalized_text = normalize_text(raw_text)

    pages: list[DocumentPage] = []
    for idx, section in enumerate(section_texts, start=1):
        pages.append(
            DocumentPage(
                page_number=idx,
                text=section,
                normalized_text=normalize_text(section),
                metadata={"section_index": idx},
            )
        )

    document_type, confidence, reasons = classify_pdf_document(
        file_path.name, normalized_text
    )

    return Document(
        source_path=file_path,
        filename=file_path.name,
        file_type=InputFileType.WORD,
        document_type=document_type,
        classification_confidence=confidence,
        classification_reasons=reasons,
        raw_text=raw_text,
        normalized_text=normalized_text,
        pages=pages,
        metadata={
            "loader": "word_loader",
            "file_size_bytes": file_path.stat().st_size,
            "source_subtype": None,
            "section_count": len(pages),
            "text_preview": collapse_snippet(normalized_text),
        },
    )
